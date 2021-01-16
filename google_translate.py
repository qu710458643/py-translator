from googletrans import Translator
from docx import Document
import re


class GoogleTranslator:
    def __init__(self):
        self.googletrans = Translator()
        return

    def GoogleTranslatText(self, text: str, src: str, dest: str) -> str:
        destPara = self.googletrans.translate(text=text, src=src, dest=dest)
        return destPara.text

    def GoogleTranslatPara(self, paragraph, src: str, dest: str, isComprision=False):
        for run in paragraph.runs:
            destText = self.GoogleTranslatText(run.text, src, dest)
            if isComprision is True:
                paragraph.add_run(destText,run.style)
            else:
                run.text = destText
            print(destText)
        return

    def GoogleTranslatFile(self, srcFileName: str, src: str, dest: str, isComprision=False) -> str:
        srcDoc = Document(srcFileName)
        splitFileName = srcFileName.split('.', 1)
        newFileName = splitFileName[0]+'_'+dest+'.'+splitFileName[1]
        for paragraph in srcDoc.paragraphs:
            if (paragraph.text == None) or (bool(re.search('[a-z]', paragraph.text)) == False):
                continue
            self.GoogleTranslatPara(paragraph, src, dest, isComprision)
        for table in srcDoc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if (paragraph.text == None) or (bool(re.search('[a-z]', paragraph.text)) == False):
                            continue
                        self.GoogleTranslatPara(paragraph, src, dest, isComprision)
        srcDoc.save(newFileName)
        return


if __name__ == "__main__":
    googleTranslator = GoogleTranslator()
    dest = googleTranslator.GoogleTranslatText("Hello", "en", "zh-cn")
    googleTranslator.GoogleTranslatFile("test.docx", "en", "zh-cn", True)
    print(dest)
