import sys
import logging
import time

from tencentcloud.common import credential
from tencentcloud.common.exception.tencent_cloud_sdk_exception import TencentCloudSDKException
# 导入对应产品模块的client models。
from tencentcloud.tmt.v20180321 import models, tmt_client

# 导入可选配置类
from tencentcloud.common.profile.client_profile import ClientProfile
from tencentcloud.common.profile.http_profile import HttpProfile
from docx import Document
from docx.shared import Inches


class MyTencentTranslator(object):
    def __init__(self):
        self.secretId = "AKIDyiDaP1VbaNxJTFrPJLBzZ9VhbJtGeBKE"
        self.secretKey = "qwnuvoYrISMG7OUQZzPRt74kCSTLsSov"
        self.client = None
        return

    def SetSecretPara(self, secretId: str, secretKey: str) -> None:
        self.secretId = secretId
        self.secretKey = secretKey
        return

    def SetHttpProxy(self, httpProfile: HttpProfile) -> None:
        # 如果需要指定proxy访问接口，可以按照如下方式初始化hp
        # httpProfile = HttpProfile(proxy="http://用户名:密码@代理IP:代理端口")
        # 在外网互通的网络环境下支持http协议(默认是https协议),建议使用https协议
        httpProfile.protocol = "https"
        httpProfile.keepAlive = True  # 状态保持，默认是False
        httpProfile.reqMethod = "GET"  # get请求(默认为post请求)
        httpProfile.reqTimeout = 30    # 请求超时时间，单位为秒(默认60秒)
        # 指定接入地域域名(默认就近接入)
        httpProfile.endpoint = "tmt.ap-shanghai.tencentcloudapi.com"
        return

    def SetClientProfile(self, clientProfile: ClientProfile, httpProfile: HttpProfile) -> None:
        clientProfile.signMethod = "TC3-HMAC-SHA256"  # 指定签名算法
        clientProfile.language = "en-US"  # 指定展示英文（默认为中文）
        clientProfile.httpProfile = httpProfile
        return

    def InitTranslator(self) -> None:
        httpProfile = HttpProfile()
        clientProfile = ClientProfile()
        self.SetHttpProxy(httpProfile)
        self.SetClientProfile(clientProfile, httpProfile)
        cred = credential.Credential(self.secretId, self.secretKey)
        self.client = tmt_client.TmtClient(cred, "ap-shanghai", clientProfile)
        return

    def TraslatorTencentPara(self, para: str) -> str:
        try:
            reqDoc = models.TextTranslateRequest()
            reqDoc.SourceText = para
            reqDoc.Source = "en"
            reqDoc.Target = "zh"
            reqDoc.ProjectId = 1
            resp = self.client.TextTranslate(reqDoc)
        except TencentCloudSDKException as err:
            print(err)
            return None
        return resp.TargetText

    def TraslatorGooglePara(self, para: str) -> str:
        
        return resp.TargetText

    def TranslateDocx(self, toolName, fileName: str, fileName2: str):
        document = Document(fileName)
        self.InitTranslator()
        # proxies={'http':'http://127.0.0.1:8889','https':'https://127.0.0.1:8889'}
        googletrans = Translator(service_urls=['translate.googleapis.com'])
        for paragraph in document.paragraphs:
            if (paragraph.text == None) or (paragraph.text == "\t"):
                continue
            if (toolName == "Google"):
                destPara = googletrans.translate(paragraph.text, dest='zh-cn')
            else:
                time.sleep(5)
                destPara = self.TraslatorPara(paragraph.text)
            paragraph.add_run('\n'+ destPara.text)
            print(destPara.text)
        document.save(fileName2)
        return


if __name__ == "__main__":
    myTencentTranslator = MyTencentTranslator()
    myTencentTranslator.TranslateDocx("Google", "test.docx", "34229-1-f60_zh.docx")
